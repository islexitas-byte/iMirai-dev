from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
import re
import os

# ========================= DEFAULT CONFIG =========================
DEFAULT_CFG = {
    "title": {"size": Pt(24), "font": "Arial Black", "color": None},
    "heading": {"size": Pt(18), "font": "Times New Roman", "color": None},
    "content": {"size": Pt(14), "font": "Calibri", "color": None},
}

# ========================= COLOR PARSER =========================
def parse_any_color(value):
    if not value:
        return None
    v = value.strip().lower()
    if v.startswith("#") and len(v) == 7:
        return RGBColor(int(v[1:3], 16), int(v[3:5], 16), int(v[5:7], 16))
    if v.startswith("rgb"):
        nums = list(map(int, re.findall(r"\d+", v)))
        if len(nums) == 3:
            return RGBColor(nums[0], nums[1], nums[2])
    return None

# ========================= STYLE PARSER =========================
def parse_styles(soup):
    cfg = {k: v.copy() for k, v in DEFAULT_CFG.items()}
    style_block = soup.select_one(".presentation-style")
    if style_block:
        for item in style_block.select("style-item"):
            dtype = item.get("data-type")
            value = item.text.strip()
            if dtype == "title-font-size":
                cfg["title"]["size"] = Pt(int(value))
            elif dtype == "title-font-style":
                cfg["title"]["font"] = value
            elif dtype == "title-color":
                cfg["title"]["color"] = parse_any_color(value)
            elif dtype == "heading-font-size":
                cfg["heading"]["size"] = Pt(int(value))
            elif dtype == "heading-font-style":
                cfg["heading"]["font"] = value
            elif dtype == "heading-color":
                cfg["heading"]["color"] = parse_any_color(value)
            elif dtype == "paragraph-font-size":
                cfg["content"]["size"] = Pt(int(value))
            elif dtype == "paragraph-font-style":
                cfg["content"]["font"] = value
            elif dtype == "paragraph-color":
                cfg["content"]["color"] = parse_any_color(value)
    return cfg

# ========================= FONT APPLY HELPER =========================
def apply_font(run, cfg_item, bold=False):
    run.font.size = cfg_item["size"]
    run.font.name = cfg_item["font"]
    run.font.bold = bold
    if cfg_item["color"]:
        run.font.color.rgb = cfg_item["color"]

# ========================= DOCUMENT BUILDERS =========================
def add_title(doc, text, cfg):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    apply_font(run, cfg["title"], bold=True)

def add_heading(doc, text, cfg):
    p = doc.add_paragraph()
    run = p.add_run(text)
    apply_font(run, cfg["heading"], bold=True)

def add_paragraph(doc, text, cfg):
    p = doc.add_paragraph()
    run = p.add_run(text)
    apply_font(run, cfg["content"])

def add_bullets(doc, items, cfg):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            apply_font(run, cfg["content"])

def add_table(doc, table_tag, cfg):
    rows = table_tag.find_all("tr")
    cols = len(rows[0].find_all(["th", "td"]))
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for r, row in enumerate(rows):
        cells = row.find_all(["th", "td"])
        for c, cell in enumerate(cells):
            p = table.rows[r].cells[c].paragraphs[0]
            run = p.add_run(cell.text.strip())
            apply_font(run, cfg["content"], bold=(cell.name == "th"))

def add_image(doc, img_path, width=None):
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(width) if width else None)

# ========================= MAIN CONVERTER WITH TEMPLATE SUPPORT =========================
def generate_docx_from_html(
    html_content, output_path, template_path=None, default_img_width=5
):
    """
    Converts HTML content to DOCX, optionally using a template.
    - html_content: HTML string with <h1>, <h2>, <p>, <ul>, <table>, <img>
    - output_path: Path to save the final DOCX
    - template_path: Optional path to existing DOCX template to append content
    - default_img_width: Default width in inches for images
    """
    soup = BeautifulSoup(html_content, "html.parser")

    # Use template if provided
    if template_path and os.path.exists(template_path):
        doc = Document(template_path)
    else:
        doc = Document()

    cfg = parse_styles(soup)

    # Track images already added to prevent duplicates
    added_images = set()

    # Title
    title_div = soup.select_one(".document-title, h1")
    if title_div:
        add_title(doc, title_div.text.strip(), cfg)

    # Sections (including paragraphs, bullets, tables, images)
    for element in soup.find_all(["h2", "p", "ul", "table", "img"]):
        if element.name == "h2":
            add_heading(doc, element.text.strip(), cfg)
        elif element.name == "p":
            img_tag = element.find("img")
            if img_tag and img_tag.has_attr("src"):
                img_src = img_tag["src"]
                if img_src not in added_images:
                    add_image(doc, img_src, width=default_img_width)
                    added_images.add(img_src)
            if element.get_text(strip=True):
                add_paragraph(doc, element.get_text(strip=True), cfg)
        elif element.name == "ul":
            items = [li.text.strip() for li in element.find_all("li")]
            add_bullets(doc, items, cfg)
        elif element.name == "table":
            add_table(doc, element, cfg)
        elif element.name == "img":
            img_src = element.get("src")
            if img_src and img_src not in added_images:
                add_image(doc, img_src, width=default_img_width)
                added_images.add(img_src)

    doc.save(output_path)
    return output_path
